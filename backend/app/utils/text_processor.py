import re
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Try to import NLTK components
try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk.stem import WordNetLemmatizer
    
    # Download required NLTK data
    for resource in ['punkt', 'stopwords', 'wordnet', 'punkt_tab']:
        try:
            nltk.data.find(f'tokenizers/{resource}' if 'punkt' in resource else f'corpora/{resource}')
        except LookupError:
            try:
                nltk.download(resource, quiet=True)
            except:
                pass
    
    STOP_WORDS = set(stopwords.words('english'))
    NLTK_AVAILABLE = True
except ImportError:
    logger.warning("NLTK not available, using basic preprocessing")
    NLTK_AVAILABLE = False
    STOP_WORDS = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                  'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
                  'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
                  'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'this',
                  'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they'}


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF file content
    
    Args:
        content: Raw PDF file bytes
        
    Returns:
        Extracted text string
    """
    text = ""
    
    # Try PyPDF2 first
    try:
        import PyPDF2
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            logger.info(f"Extracted {len(text)} characters using PyPDF2")
            return clean_extracted_text(text)
    except ImportError:
        logger.warning("PyPDF2 not installed")
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # Try pdfplumber as fallback
    try:
        import pdfplumber
        pdf_file = io.BytesIO(content)
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            logger.info(f"Extracted {len(text)} characters using pdfplumber")
            return clean_extracted_text(text)
    except ImportError:
        logger.warning("pdfplumber not installed")
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Try PyMuPDF (fitz) as another fallback
    try:
        import fitz  # PyMuPDF
        pdf_file = io.BytesIO(content)
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        
        for page in doc:
            text += page.get_text() + "\n"
        
        doc.close()
        
        if text.strip():
            logger.info(f"Extracted {len(text)} characters using PyMuPDF")
            return clean_extracted_text(text)
    except ImportError:
        logger.warning("PyMuPDF not installed")
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")
    
    logger.error("All PDF extraction methods failed")
    return ""


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX file content
    
    Args:
        content: Raw DOCX file bytes
        
    Returns:
        Extracted text string
    """
    text = ""
    
    # Try python-docx
    try:
        from docx import Document
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        
        if text.strip():
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return clean_extracted_text(text)
    except ImportError:
        logger.warning("python-docx not installed")
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
    
    # Fallback: Try to decode as text (for corrupted files)
    try:
        text = content.decode('utf-8', errors='ignore')
        # Extract readable portions
        text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    except:
        pass
    
    return ""


def clean_extracted_text(text: str) -> str:
    """
    Clean extracted text from PDF/DOCX
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove page numbers
    text = re.sub(r'\bPage\s*\d+\s*(?:of\s*\d+)?\b', '', text, flags=re.IGNORECASE)
    
    # Remove common PDF artifacts
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # Fix common encoding issues
    text = text.replace('â€™', "'").replace('â€"', "-").replace('â€œ', '"').replace('â€', '"')
    
    # Remove URLs (keep for reference but clean them)
    text = re.sub(r'https?://\S+', '[URL]', text)
    
    # Clean up spacing around punctuation
    text = re.sub(r'\s+([.,;:!?])', r'\1', text)
    text = re.sub(r'([.,;:!?])\s+', r'\1 ', text)
    
    return text.strip()


def preprocess_text(text: str) -> str:
    """
    Preprocess text for ML model input
    
    Args:
        text: Raw text to preprocess
        
    Returns:
        Preprocessed text suitable for ML model
    """
    if not text:
        return ""
    
    # Step 1: Lowercase
    text = text.lower()
    
    # Step 2: Remove special characters but keep spaces
    text = re.sub(r'[^a-zA-Z\s]', ' ', text)
    
    # Step 3: Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Step 4: Tokenization and stopword removal
    if NLTK_AVAILABLE:
        try:
            tokens = word_tokenize(text)
            lemmatizer = WordNetLemmatizer()
            
            clean_tokens = [
                lemmatizer.lemmatize(token)
                for token in tokens
                if token not in STOP_WORDS and len(token) > 2
            ]
            
            return " ".join(clean_tokens)
        except Exception as e:
            logger.warning(f"NLTK processing failed: {e}, using basic preprocessing")
    
    # Fallback: Basic preprocessing without NLTK
    words = text.split()
    clean_words = [
        word for word in words
        if word not in STOP_WORDS and len(word) > 2
    ]
    
    return " ".join(clean_words)


def extract_contact_info(text: str) -> dict:
    """
    Extract contact information from resume text
    
    Args:
        text: Resume text
        
    Returns:
        Dictionary with extracted contact info
    """
    contact = {
        "email": None,
        "phone": None,
        "linkedin": None,
        "github": None,
        "website": None
    }
    
    # Email extraction
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact["email"] = email_match.group()
    
    # Phone extraction
    phone_patterns = [
        r'\+?1?\s*$$?[0-9]{3}$$?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
        r'\+?[0-9]{1,3}[-.\s]?[0-9]{10}',
        r'[0-9]{10}'
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            contact["phone"] = phone_match.group()
            break
    
    # LinkedIn extraction
    linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9-]+)'
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"
    
    # GitHub extraction
    github_pattern = r'github\.com/([a-zA-Z0-9-]+)'
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    if github_match:
        contact["github"] = f"github.com/{github_match.group(1)}"
    
    # Website extraction
    website_pattern = r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z]{2,})(?:/\S*)?'
    website_matches = re.findall(website_pattern, text)
    for match in website_matches:
        if not any(x in match.lower() for x in ['linkedin', 'github', 'gmail', 'yahoo', 'outlook']):
            contact["website"] = match
            break
    
    return contact


def extract_sections(text: str) -> dict:
    """
    Extract different sections from resume
    
    Args:
        text: Resume text
        
    Returns:
        Dictionary with identified sections
    """
    sections = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "certifications": "",
        "other": ""
    }
    
    # Section headers to look for
    section_patterns = {
        "summary": r'(?:summary|objective|profile|about\s*me)',
        "experience": r'(?:experience|work\s*history|employment|professional\s*experience)',
        "education": r'(?:education|academic|qualifications|degrees)',
        "skills": r'(?:skills|technical\s*skills|competencies|expertise)',
        "projects": r'(?:projects|portfolio|work\s*samples)',
        "certifications": r'(?:certifications|certificates|licenses|credentials)'
    }
    
    text_lower = text.lower()
    
    # Find section positions
    section_positions = []
    for section_name, pattern in section_patterns.items():
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            section_positions.append((match.start(), section_name))
    
    # Sort by position
    section_positions.sort(key=lambda x: x[0])
    
    # Extract section content
    for i, (pos, section_name) in enumerate(section_positions):
        end_pos = section_positions[i + 1][0] if i + 1 < len(section_positions) else len(text)
        sections[section_name] = text[pos:end_pos].strip()
    
    return sections


def count_keywords(text: str, keywords: list) -> dict:
    """
    Count occurrences of keywords in text
    
    Args:
        text: Text to search
        keywords: List of keywords to count
        
    Returns:
        Dictionary with keyword counts
    """
    text_lower = text.lower()
    counts = {}
    
    for keyword in keywords:
        keyword_lower = keyword.lower()
        # Use word boundary matching
        pattern = r'\b' + re.escape(keyword_lower) + r'\b'
        count = len(re.findall(pattern, text_lower))
        if count > 0:
            counts[keyword] = count
    
    return counts


def calculate_text_stats(text: str) -> dict:
    """
    Calculate various statistics about the text
    
    Args:
        text: Text to analyze
        
    Returns:
        Dictionary with text statistics
    """
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    
    return {
        "word_count": len(words),
        "sentence_count": len([s for s in sentences if s.strip()]),
        "avg_word_length": sum(len(w) for w in words) / max(len(words), 1),
        "avg_sentence_length": len(words) / max(len([s for s in sentences if s.strip()]), 1),
        "unique_words": len(set(w.lower() for w in words)),
        "character_count": len(text)
    }